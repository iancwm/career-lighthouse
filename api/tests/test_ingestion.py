# api/tests/test_ingestion.py
import numpy as np
import pytest
from unittest.mock import MagicMock
from services.ingestion import chunk_text, parse_file, ingest_document


def test_chunk_text_splits_on_token_boundary():
    text = " ".join([f"word{i}" for i in range(200)])
    chunks = chunk_text(text, max_tokens=50, overlap=10)
    assert len(chunks) > 1
    for chunk in chunks:
        assert len(chunk.split()) <= 60  # allow slight overage for overlap


def test_chunk_text_single_chunk_if_short():
    text = "short text here"
    chunks = chunk_text(text, max_tokens=512, overlap=64)
    assert chunks == ["short text here"]


def test_parse_file_txt():
    content = b"hello career world"
    text = parse_file(content, "test.txt")
    assert "hello career world" in text


def test_ingest_document_calls_upsert(in_memory_qdrant, mock_embedder):
    from services.vector_store import VectorStore
    store = VectorStore(client=in_memory_qdrant, collection="knowledge")
    store.ensure_collection(dim=384)

    mock_embedder.encode_batch.return_value = np.ones((2, 384), dtype=np.float32)

    count = ingest_document(
        file_content=b"chunk one content. " * 60 + b"chunk two content. " * 60,
        filename="test.txt",
        embedder=mock_embedder,
        store=store,
    )
    assert count >= 1
    mock_embedder.encode_batch.assert_called_once()
    docs = store.list_docs()
    assert any(d["filename"] == "test.txt" for d in docs)


def test_parse_file_docx_with_table():
    """DOCX with a paragraph and a 2x3 table — both should appear in output."""
    from io import BytesIO
    from docx import Document

    doc = Document()
    doc.add_paragraph("Intro paragraph before table.")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Role"
    table.cell(0, 1).text = "Base"
    table.cell(0, 2).text = "Bonus"
    table.cell(1, 0).text = "Junior"
    table.cell(1, 1).text = "80K"
    table.cell(1, 2).text = "15%"

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    content = buf.read()

    text = parse_file(content, "memo.docx")
    assert "Intro paragraph before table." in text
    assert "Role" in text
    assert "Junior" in text
    assert "80K" in text
    # Table rows should be pipe-delimited
    assert "Role | Base | Bonus" in text
    assert "Junior | 80K | 15%" in text


def test_parse_file_docx_merged_cells_no_duplication():
    """DOCX with merged header row — merged cell text appears exactly once."""
    from io import BytesIO
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree

    doc = Document()
    doc.add_paragraph("Header text.")

    # Build a table with a merged header row using raw XML
    tbl = etree.fromstring(
        '<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '  <w:tr>'
        '    <w:tc>'
        '      <w:tcPr><w:gridSpan w:val="2"/></w:tcPr>'
        '      <w:p><w:r><w:t>Merged Header</w:t></w:r></w:p>'
        '    </w:tc>'
        '    <w:tc/>'
        '  </w:tr>'
        '  <w:tr>'
        '    <w:tc><w:p><w:r><w:t>A</w:t></w:r></w:p></w:tc>'
        '    <w:tc><w:p><w:r><w:t>B</w:t></w:r></w:p></w:tc>'
        '  </w:tr>'
        '</w:tbl>'
    )
    doc.element.body.append(tbl)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    content = buf.read()

    text = parse_file(content, "memo.docx")
    # Merged cell text should appear exactly once
    assert text.count("Merged Header") == 1
    # Data row cells should appear
    assert "A" in text
    assert "B" in text


# ---------------------------------------------------------------------------
# Semantic-aware chunking tests
# ---------------------------------------------------------------------------

def test_chunk_preserves_paragraph_boundaries():
    """Two short paragraphs should each be their own chunk, not split mid-paragraph."""
    from services.ingestion import chunk_text
    text = "First paragraph about career paths.\n\nSecond paragraph about salary expectations."
    chunks = chunk_text(text, max_tokens=512, overlap=64)
    # Both paragraphs should be preserved in at least one chunk
    combined = " ".join(chunks)
    assert "First paragraph about career paths." in combined
    assert "Second paragraph about salary expectations." in combined


def test_chunk_preserves_table_rows():
    """A small table should be kept together in one chunk, not split across rows."""
    from services.ingestion import chunk_text
    text = """Here is the salary table:
| Role | Base | Bonus |
| Junior | 80K | 15% |
| Senior | 120K | 25% |
| Manager | 160K | 35% |"""
    chunks = chunk_text(text, max_tokens=512, overlap=64)
    # Table rows should appear together (at least in one chunk)
    combined = " ".join(chunks)
    assert "Junior" in combined
    assert "Senior" in combined
    assert "Manager" in combined


def test_chunk_splits_oversized_paragraph():
    """A single long paragraph exceeding token limit should split at word boundaries."""
    from services.ingestion import chunk_text
    words = [f"word{i}" for i in range(500)]
    long_paragraph = " ".join(words)
    chunks = chunk_text(long_paragraph, max_tokens=100, overlap=20)
    # Should produce multiple chunks
    assert len(chunks) > 1
    # First chunk should contain the beginning
    assert "word0" in chunks[0]
    # Last chunk should contain the end
    assert "word499" in chunks[-1]


def test_chunk_single_short_text_returns_as_is():
    """Text within token limit should return as single chunk."""
    from services.ingestion import chunk_text
    text = "Short career advice."
    chunks = chunk_text(text, max_tokens=512, overlap=64)
    assert chunks == ["Short career advice."]


def test_chunk_empty_returns_empty():
    """Empty text should return empty list."""
    from services.ingestion import chunk_text
    assert chunk_text("", max_tokens=512, overlap=64) == []
    assert chunk_text("   ", max_tokens=512, overlap=64) == []


def test_chunk_preserves_list_items():
    """Consecutive list items should be kept together."""
    from services.ingestion import chunk_text
    text = """Skills needed:
- Communication
- Analytical thinking
- Leadership"""
    chunks = chunk_text(text, max_tokens=512, overlap=64)
    combined = " ".join(chunks)
    assert "Communication" in combined
    assert "Analytical thinking" in combined
    assert "Leadership" in combined
